from __future__ import annotations

import html
import ipaddress
import re
import unicodedata
from dataclasses import dataclass, field
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from urllib.parse import parse_qsl, urlencode, urlsplit, urlunsplit


DOCUMENT_EXPORT_SCHEMA_VERSION = 1
_MARKDOWN_LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
_BULLET_RE = re.compile(r"^\s*[-*+]\s+(.+?)\s*$")
_ORDERED_RE = re.compile(r"^\s*\d+[.)]\s+(.+?)\s*$")
_RAW_URL_RE = re.compile(r"https?://[^\s<>\]\[\"']+", re.IGNORECASE)
_SECRET_ASSIGNMENT_RE = re.compile(
    r"(?im)\b(cookie|set-cookie|authorization|proxy-authorization|password|secret|"
    r"access[_-]?token|refresh[_-]?token|auth[_-]?token|session(?:id)?|api[_-]?key)"
    r"\b(\s*[:=]\s*)[^\r\n]*"
)
_BEARER_RE = re.compile(r"(?i)\b(Bearer|Basic)\s+[A-Za-z0-9._~+/=-]+")
_SAFE_URL_QUERY_KEYS = {"v", "p", "list", "index", "t", "start", "page", "bvid", "aid", "cid"}


class DocumentExportUnavailable(RuntimeError):
    """Raised when an explicitly declared document dependency is unavailable."""


@dataclass(frozen=True)
class DocumentExport:
    content: bytes
    media_type: str
    suffix: str
    font_name: str
    warnings: list[str] = field(default_factory=list)
    schema_version: int = DOCUMENT_EXPORT_SCHEMA_VERSION


@dataclass(frozen=True)
class _Block:
    kind: str
    text: str
    level: int = 0


def _blocks(markdown: str) -> list[_Block]:
    """Parse a conservative Markdown subset while preserving unknown content."""
    result: list[_Block] = []
    paragraph: list[str] = []
    code: list[str] = []
    in_code = False

    def flush_paragraph() -> None:
        if paragraph:
            result.append(_Block("paragraph", " ".join(item.strip() for item in paragraph if item.strip())))
            paragraph.clear()

    def flush_code() -> None:
        if code:
            result.append(_Block("code", "\n".join(code)))
            code.clear()

    for raw_line in str(markdown or "").replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        line = raw_line.rstrip()
        if line.strip().startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                flush_paragraph()
                in_code = True
            continue
        if in_code:
            code.append(line)
            continue
        if not line.strip():
            flush_paragraph()
            continue
        if re.fullmatch(r"\s*([-*_])(?:\s*\1){2,}\s*", line):
            flush_paragraph()
            continue
        heading = _HEADING_RE.match(line)
        if heading:
            flush_paragraph()
            result.append(_Block("heading", heading.group(2).strip(), len(heading.group(1))))
            continue
        bullet = _BULLET_RE.match(line)
        if bullet:
            flush_paragraph()
            result.append(_Block("bullet", bullet.group(1).strip()))
            continue
        ordered = _ORDERED_RE.match(line)
        if ordered:
            flush_paragraph()
            result.append(_Block("ordered", ordered.group(1).strip()))
            continue
        paragraph.append(line)
    flush_paragraph()
    flush_code()
    return [block for block in result if block.text.strip()]


def _safe_hyperlink(value: str) -> str:
    target = str(value or "").strip()
    try:
        parsed = urlsplit(target)
    except ValueError:
        return ""
    try:
        if parsed.scheme.lower() not in {"http", "https"} or not parsed.netloc or parsed.username or parsed.password:
            return ""
        hostname = parsed.hostname or ""
        lowered = hostname.rstrip(".").lower()
        if lowered in {"localhost", "localhost.localdomain"} or lowered.endswith((".localhost", ".local", ".lan", ".internal", ".home", ".corp")):
            return ""
        try:
            if not ipaddress.ip_address(lowered).is_global:
                return ""
        except ValueError:
            pass
        port = f":{parsed.port}" if parsed.port else ""
        query = [
            (str(key)[:40], str(item)[:200])
            for key, item in parse_qsl(parsed.query, keep_blank_values=True, max_num_fields=100)
            if key.lower() in _SAFE_URL_QUERY_KEYS
        ]
        fragment = parsed.fragment if parsed.fragment.lower().startswith(("t=", "page=")) else ""
        return urlunsplit((parsed.scheme.lower(), hostname + port, parsed.path[:1200], urlencode(query), fragment[:120]))[:1600]
    except (TypeError, ValueError):
        return ""


def sanitize_export_text(value: str) -> str:
    text = unicodedata.normalize("NFC", str(value or "").replace("\r\n", "\n").replace("\r", "\n"))
    text = "".join(character for character in text if character in "\n\t" or ord(character) >= 0x20)
    text = _SECRET_ASSIGNMENT_RE.sub(lambda match: f"{match.group(1)}{match.group(2)}[REDACTED]", text)
    text = _BEARER_RE.sub(lambda match: f"{match.group(1)} [REDACTED]", text)

    def replace_url(match: re.Match[str]) -> str:
        safe = _safe_hyperlink(match.group(0).rstrip(".,;，。；"))
        return safe or "[private URL removed]"

    return _RAW_URL_RE.sub(replace_url, text)


_sanitize_export_text = sanitize_export_text


def _clean_inline_markdown(value: str) -> str:
    text = str(value or "")
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"__([^_]+)__", r"\1", text)
    return text


def _content_blocks(markdown: str, title: str) -> list[_Block]:
    blocks = _blocks(markdown)
    if not blocks:
        return blocks
    first = blocks[0]
    def normalize(value: str) -> str:
        return re.sub(r"\s+", " ", _clean_inline_markdown(value)).strip().casefold()

    if first.kind == "heading" and first.level == 1 and normalize(first.text) == normalize(title):
        return blocks[1:]
    return blocks


def _add_docx_hyperlink(paragraph, text: str, url: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    relation_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0F766E")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend((color, underline))
    run.append(properties)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_docx_inline(paragraph, value: str) -> None:
    cursor = 0
    text = _clean_inline_markdown(value)
    for match in _MARKDOWN_LINK_RE.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(_sanitize_export_text(text[cursor:match.start()]))
        label, raw_url = match.group(1), match.group(2)
        url = _safe_hyperlink(raw_url)
        if url:
            display_label = url if "://" in label else _sanitize_export_text(label)
            _add_docx_hyperlink(paragraph, display_label, url)
        else:
            paragraph.add_run(f"{_sanitize_export_text(label)}（链接已移除）")
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(_sanitize_export_text(text[cursor:]))


def build_docx_export(task, note: str, transcript: dict | None = None) -> DocumentExport:
    try:
        from docx import Document
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt
    except ImportError as exc:
        raise DocumentExportUnavailable("docx_export_dependency_missing") from exc

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.9)
    section.bottom_margin = Cm(1.9)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)
    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    for name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        style = document.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    raw_title = str(getattr(task, "title", "LearnNote 学习笔记")) or "LearnNote 学习笔记"
    safe_title = _sanitize_export_text(raw_title)[:250]
    document.core_properties.title = safe_title
    document.core_properties.subject = "LearnNote evidence-grounded local export"
    document.core_properties.author = "LearnNote"
    document.add_heading(safe_title, 0)

    source_url = _safe_hyperlink(str(getattr(task, "page_url", "")))
    source_line = document.add_paragraph()
    source_line.add_run("来源：").bold = True
    if source_url:
        _add_docx_hyperlink(source_line, source_url, source_url)
    else:
        source_line.add_run("本地资料")
    source_line.add_run(f"\n导出时间：{datetime.now(timezone.utc).isoformat()}")
    segments = (transcript or {}).get("segments") if isinstance(transcript, dict) else []
    if isinstance(segments, list) and segments:
        source_line.add_run(f"\n可追溯字幕片段：{len(segments)} 段")

    for block in _content_blocks(note, raw_title):
        if block.kind == "heading":
            paragraph = document.add_heading(level=max(1, min(block.level, 3)))
            _add_docx_inline(paragraph, block.text)
        elif block.kind == "bullet":
            paragraph = document.add_paragraph(style="List Bullet")
            _add_docx_inline(paragraph, block.text)
        elif block.kind == "ordered":
            paragraph = document.add_paragraph(style="List Number")
            _add_docx_inline(paragraph, block.text)
        elif block.kind == "code":
            paragraph = document.add_paragraph()
            run = paragraph.add_run(_sanitize_export_text(block.text))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        else:
            paragraph = document.add_paragraph()
            _add_docx_inline(paragraph, block.text)

    # Word handles page breaks and long documents more reliably when the final
    # paragraph is not part of a list or a code run.
    document.add_paragraph("由 LearnNote 在本机生成；原视频、Cookie 与诊断秘密未嵌入此文档。")
    buffer = BytesIO()
    document.save(buffer)
    return DocumentExport(
        content=buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        suffix="docx",
        font_name="Microsoft YaHei (document preference with host fallback)",
    )


def _pdf_font() -> tuple[str, list[str]]:
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.pdfbase.ttfonts import TTFont

    candidates = (
        ("LearnNoteCJK", Path("C:/Windows/Fonts/msyh.ttc")),
        ("LearnNoteCJK", Path("C:/Windows/Fonts/simsun.ttc")),
        ("LearnNoteCJK", Path("/System/Library/Fonts/PingFang.ttc")),
        ("LearnNoteCJK", Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc")),
        ("LearnNoteCJK", Path("/usr/share/fonts/truetype/wqy/wqy-microhei.ttc")),
    )
    for name, path in candidates:
        if not path.is_file():
            continue
        try:
            pdfmetrics.registerFont(TTFont(name, str(path), subfontIndex=0))
            return name, []
        except Exception:
            continue
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    return "STSong-Light", ["embedded_system_cjk_font_unavailable_using_pdf_cid_fallback"]


def _pdf_inline(value: str) -> str:
    source = _clean_inline_markdown(value)
    parts: list[str] = []
    cursor = 0
    for match in _MARKDOWN_LINK_RE.finditer(source):
        parts.append(html.escape(_sanitize_export_text(source[cursor:match.start()])))
        label, raw_url = match.group(1), match.group(2)
        url = _safe_hyperlink(raw_url)
        if url:
            display_label = url if "://" in label else _sanitize_export_text(label)
            parts.append(f'<a href="{html.escape(url, quote=True)}" color="#0f766e"><u>{html.escape(display_label)}</u></a>')
        else:
            parts.append(html.escape(f"{_sanitize_export_text(label)}（链接已移除）"))
        cursor = match.end()
    parts.append(html.escape(_sanitize_export_text(source[cursor:])))
    return "".join(parts)


def build_pdf_export(task, note: str, transcript: dict | None = None) -> DocumentExport:
    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.platypus import Paragraph, Preformatted, SimpleDocTemplate, Spacer
    except ImportError as exc:
        raise DocumentExportUnavailable("pdf_export_dependency_missing") from exc

    font_name, warnings = _pdf_font()
    buffer = BytesIO()
    document = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
        title=_sanitize_export_text(str(getattr(task, "title", "LearnNote 学习笔记"))),
        author="LearnNote",
    )
    sample = getSampleStyleSheet()
    body = ParagraphStyle(
        "LearnNoteBody",
        parent=sample["BodyText"],
        fontName=font_name,
        fontSize=10.5,
        leading=17,
        textColor=colors.HexColor("#17201F"),
        spaceAfter=7,
        wordWrap="CJK",
    )
    title_style = ParagraphStyle(
        "LearnNoteTitle",
        parent=body,
        fontSize=21,
        leading=28,
        textColor=colors.HexColor("#123B37"),
        spaceAfter=12,
    )
    heading_styles = {
        1: ParagraphStyle("LearnNoteH1", parent=body, fontSize=16, leading=23, textColor=colors.HexColor("#0F5F58"), spaceBefore=12, spaceAfter=7),
        2: ParagraphStyle("LearnNoteH2", parent=body, fontSize=13.5, leading=20, textColor=colors.HexColor("#155E58"), spaceBefore=10, spaceAfter=6),
        3: ParagraphStyle("LearnNoteH3", parent=body, fontSize=11.5, leading=18, textColor=colors.HexColor("#1D514D"), spaceBefore=8, spaceAfter=4),
    }
    meta = ParagraphStyle("LearnNoteMeta", parent=body, fontSize=8.5, leading=13, textColor=colors.HexColor("#52615F"))
    code = ParagraphStyle("LearnNoteCode", parent=body, fontName=font_name, fontSize=8.5, leading=12, backColor=colors.HexColor("#F2F6F5"), borderPadding=7)
    bullet = ParagraphStyle("LearnNoteBullet", parent=body, leftIndent=10, firstLineIndent=-8, bulletIndent=0)

    raw_title = str(getattr(task, "title", "LearnNote 学习笔记")) or "LearnNote 学习笔记"
    safe_title = _sanitize_export_text(raw_title)
    story = [Paragraph(html.escape(safe_title), title_style)]
    source_url = _safe_hyperlink(str(getattr(task, "page_url", "")))
    source = f'<a href="{html.escape(source_url, quote=True)}" color="#0f766e">{html.escape(source_url)}</a>' if source_url else "本地资料"
    segments = (transcript or {}).get("segments") if isinstance(transcript, dict) else []
    segment_text = f" · 可追溯字幕片段 {len(segments)} 段" if isinstance(segments, list) and segments else ""
    story.extend((
        Paragraph(f"来源：{source}<br/>导出时间：{datetime.now(timezone.utc).isoformat()}{segment_text}", meta),
        Spacer(1, 5 * mm),
    ))
    for block in _content_blocks(note, raw_title):
        if block.kind == "heading":
            story.append(Paragraph(_pdf_inline(block.text), heading_styles[max(1, min(block.level, 3))]))
        elif block.kind == "bullet":
            story.append(Paragraph(_pdf_inline(block.text), bullet, bulletText="•"))
        elif block.kind == "ordered":
            story.append(Paragraph("• " + _pdf_inline(block.text), bullet))
        elif block.kind == "code":
            story.append(Preformatted(_sanitize_export_text(block.text), code))
        else:
            story.append(Paragraph(_pdf_inline(block.text), body))
    story.extend((Spacer(1, 4 * mm), Paragraph("由 LearnNote 在本机生成；原视频、Cookie 与诊断秘密未嵌入此文档。", meta)))

    def draw_footer(canvas, doc) -> None:
        canvas.saveState()
        canvas.setFont(font_name, 8)
        canvas.setFillColor(colors.HexColor("#71807E"))
        canvas.drawCentredString(A4[0] / 2, 9 * mm, f"LearnNote · {doc.page}")
        canvas.restoreState()

    document.build(story, onFirstPage=draw_footer, onLaterPages=draw_footer)
    return DocumentExport(
        content=buffer.getvalue(),
        media_type="application/pdf",
        suffix="pdf",
        font_name=font_name,
        warnings=warnings,
    )
